import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE  # Necessario per evitare l'errore "KeyError: no style"
from docx.oxml import OxmlElement, ns
import io
import fitz
from openai import OpenAI

st.set_page_config(page_title="KDP Professional Tool", layout="wide")

# --- RIGHE AGGIUNTE: NASCONDE IL MENU IN ALTO A DESTRA E IL FOOTER ---
hide_streamlit_style = """
<style>
#MainMenu {visibility: hidden;}
header {visibility: hidden;}
footer {visibility: hidden;}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)
# ---------------------------------------------------------------------

try:
    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
except Exception:
    st.error("Configura OPENAI_API_KEY nei Secrets di Streamlit.")
    st.stop()

# --- LOGICA DI PULIZIA E FORMATTAZIONE ---

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        run = p.add_run()
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

def impeccable_format(file):
    doc = Document(file)
    
    # Crea gli stili se il documento originale ne è sprovvisto (Previene il crash dell'app)
    for style_name in ['Heading 1', 'Heading 2']:
        try:
            doc.styles[style_name]
        except KeyError:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    for section in doc.sections:
        section.page_width = Inches(6)
        section.page_height = Inches(9)
        # --- MARGINI SIMMETRICI (Tutti i lati uguali) ---
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        # ------------------------------------------------

    for p in list(doc.paragraphs):
        # Salva l'interlinea originale del documento
        spazio_originale = p.paragraph_format.line_spacing
        regola_spazio_originale = p.paragraph_format.line_spacing_rule

        text = p.text.strip()
        if not text:
            delete_paragraph(p)
            continue
        clean_text = text.replace("**", "").replace("##", "").replace("#", "")
        p.text = " ".join(clean_text.split())
        
        if len(p.text) < 60 and ("CAPITOLO" in p.text.upper() or p.text.isupper()):
            p.style = 'Heading 1'
            p.paragraph_format.page_break_before = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(30)
            
        elif len(p.text) < 80 and p.text[0].isdigit() and " " in p.text:
            p.style = 'Heading 2'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(14)
        
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.25)
            
            # Riapplica lo spazio tra le righe originale
            if spazio_originale is not None:
                p.paragraph_format.line_spacing = spazio_originale
            if regola_spazio_originale is not None:
                p.paragraph_format.line_spacing_rule = regola_spazio_originale
            
            p.paragraph_format.space_after = Pt(6)

    for style in doc.styles:
        if 'TOC' in style.name: style.font.size = Pt(10)
        if style.name == 'Normal':
            style.font.name = 'Georgia'
            style.font.size = Pt(11)

    add_page_numbers(doc)
    out_buffer = io.BytesIO()
    doc.save(out_buffer)
    out_buffer.seek(0)
    return out_buffer

# --- INTERFACCIA ---
st.title("FORMATTA IL TUO LIBRO !")

uploaded_file = st.file_uploader("Carica Manoscritto", type=["docx", "pdf"])

if uploaded_file:
    file_ext = uploaded_file.name.split(".")[-1].lower()
    col1, col2 = st.columns(2)

    with col1:
        st.header("✍️ Generazione Metadati")
        
        # --- RIGHE AGGIUNTE: SCELTA DELLA LINGUA ---
        lingue_disponibili = ["Italiano", "Inglese", "Spagnolo", "Francese", "Tedesco", "Rumeno", "Russo", "Arabo", "Cinese"]
        lingua_scelta = st.selectbox("Seleziona la lingua di traduzione/generazione:", lingue_disponibili)
        # -------------------------------------------

        if st.button("Genera Metadati Dettagliati"):
            with st.spinner(f"Generazione in {lingua_scelta} in corso..."):
                context = ""
                if file_ext == "docx":
                    d = Document(uploaded_file)
                    context = "\n".join([p.text for p in d.paragraphs[:100]])
                else:
                    with fitz.open(stream=uploaded_file.read(), filetype="pdf") as pdf:
                        for page in pdf[:15]: context += page.get_text()
                uploaded_file.seek(0)

                # --- RIGHE MODIFICATE: LINGUA DINAMICA E 7 PAROLE CHIAVE ---
                prompt = f"""
                Analizza questo libro: {context[:8000]}
                
                Genera esclusivamente il contenuto finale in lingua {lingua_scelta.upper()} e in TESTO SEMPLICE. 
                NON aggiungere introduzioni, commenti, descrizioni del tuo ragionamento o spiegazioni. 
                L'output deve contenere SOLO:

                1. DESCRIZIONE MARKETING (Dettagliata, 450+ parole): 
                   - Un gancio (hook) iniziale potente.
                   - Analisi del problema e della soluzione offerta dal libro.
                   - Elenco puntato (usa '-') dei benefici e di cosa imparerà il lettore.
                   - Call to action finale.

                2. 7 KEYWORD A CODA LUNGA: 
                   - Fornisci solo l'elenco delle 7 frasi specifiche separate da virgola.
                """
                # ----------------------------------------------------------
                
                resp = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[{"role": "system", "content": "Sei un generatore di metadati KDP. Fornisci solo il risultato finale senza commenti."},
                              {"role": "user", "content": prompt}]
                )
                st.text_area("Copia da qui:", value=resp.choices[0].message.content, height=500)

    with col2:
        st.header("⚙️ Formattazione 6x9")
        if file_ext == "docx":
            if st.button("Formatta Documento"):
                with st.spinner("Elaborazione..."):
                    clean_file = impeccable_format(uploaded_file)
                    st.success("✅ Formattazione completata!")
                    st.download_button("⬇️ Scarica Word 6x9", clean_file, f"KDP_FINAL_{uploaded_file.name}")
